import warnings

from PyQt5.QtWidgets import QVBoxLayout, QWidget

warnings.filterwarnings("ignore", category=DeprecationWarning)

import sys
import os
import math  # Add this import
from datetime import date
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtPrintSupport import QPrinter, QPrintDialog, QPrintPreviewDialog
from docx import Document
from docx2pdf import convert
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QPushButton
import fitz  # PyMuPDF
from PyQt5.QtWidgets import QLabel, QScrollArea
from PyQt5.QtGui import QPixmap, QImage
from PyQt5.QtWidgets import QHBoxLayout
from PyQt5.QtCore import QThreadPool
from backend.adminBack import adminPageBack

from PyQt5.QtCore import QRunnable, QThreadPool, QObject, pyqtSignal

from PyQt5.QtCore import QObject, QRunnable, pyqtSignal, pyqtSlot

class BillWorkerSignals(QObject):
    progress = pyqtSignal(int, int)  # i, total
    finished = pyqtSignal(str)       # merged_pdf_path
    error = pyqtSignal(str)

class BillWorker(QRunnable):
    def __init__(self, data, template_path, temp_folder, fill_word_template, convert_to_pdf):
        super().__init__()
        self.signals = BillWorkerSignals()
        self.data = data
        self.template_path = template_path
        self.temp_folder = temp_folder
        self.fill_word_template = fill_word_template
        self.convert_to_pdf = convert_to_pdf

    @pyqtSlot()
    def run(self):
        from PyPDF2 import PdfMerger
        import pythoncom
        import os

        pythoncom.CoInitialize()

        pdf_paths = []
        backend = adminPageBack()

        try:
            for i, billing in enumerate(self.data):
                billing_code = billing[0]
                billing_id = backend.get_billing_id(billing_code)
                bill_data = backend.get_bill_data_by_code(billing_id)
                if not bill_data:
                    continue

                docx_path = os.path.join(self.temp_folder, f"bill_{i}.docx")
                pdf_path = os.path.join(self.temp_folder, f"bill_{i}.pdf")

                self.fill_word_template(bill_data, self.template_path, docx_path)
                self.convert_to_pdf(docx_path, pdf_path)
                pdf_paths.append(pdf_path)

                self.signals.progress.emit(i + 1, len(self.data))

            merged_pdf_path = os.path.join(self.temp_folder, "merged_batch_print.pdf")
            merger = PdfMerger()
            for path in pdf_paths:
                merger.append(path)
            merger.write(merged_pdf_path)
            merger.close()

            self.signals.finished.emit(merged_pdf_path)

        except Exception as e:
            self.signals.error.emit(str(e))

        finally:
            pythoncom.CoUninitialize()

class QuickIssueDialog(QtWidgets.QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Quick Issue Billing")
        self.setMinimumSize(800, 500)
        self.billing_data = []
        self.filtered_data = []
        self.parent_window = parent

        self.setup_ui()
        self.load_billing_data()

    def setup_ui(self):
        layout = QtWidgets.QVBoxLayout(self)

        # Search Bar
        search_layout = QtWidgets.QHBoxLayout()
        self.search_input = QtWidgets.QLineEdit()
        self.search_input.setPlaceholderText("Search by Billing Code or Client Name...")
        self.search_input.textChanged.connect(self.filter_table)
        search_layout.addWidget(QtWidgets.QLabel("Search:"))
        search_layout.addWidget(self.search_input)

        layout.addLayout(search_layout)

        # Table Setup
        self.table = QtWidgets.QTableWidget()
        self.table.setColumnCount(3)
        self.table.setHorizontalHeaderLabels(["Select", "Billing Code", "Client Name"])
        self.table.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.Stretch)
        self.table.setSortingEnabled(True)
        layout.addWidget(self.table)

        # Button Row
        btn_layout = QtWidgets.QHBoxLayout()
        self.issue_btn = QtWidgets.QPushButton("✅ Issue Selected Bills")
        self.issue_btn.clicked.connect(self.on_issue_clicked)
        btn_layout.addStretch()
        btn_layout.addWidget(self.issue_btn)

        layout.addLayout(btn_layout)

    def load_billing_data(self):
        backend = adminPageBack()
        raw_data = backend.fetch_billing_to_issue()  # You'll define this function
        self.billing_data = raw_data
        self.filtered_data = raw_data
        self.update_table()

    def filter_table(self):
        query = self.search_input.text().strip().lower()
        if not query:
            self.filtered_data = self.billing_data
        else:
            self.filtered_data = [
                row for row in self.billing_data
                if query in row[0].lower() or query in row[2].lower()
            ]
        self.update_table()

    def update_table(self):
        self.table.setRowCount(len(self.filtered_data))
        for i, (billing_code, issued_date, client_name) in enumerate(self.filtered_data):
            checkbox = QtWidgets.QTableWidgetItem()
            checkbox.setFlags(QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
            checkbox.setCheckState(QtCore.Qt.Unchecked)

            self.table.setItem(i, 0, checkbox)
            self.table.setItem(i, 1, QtWidgets.QTableWidgetItem(billing_code))
            self.table.setItem(i, 2, QtWidgets.QTableWidgetItem(client_name))

    def on_issue_clicked(self):
        selected_rows = []
        for i in range(self.table.rowCount()):
            item = self.table.item(i, 0)
            if item.checkState() == QtCore.Qt.Checked:
                billing_code = self.table.item(i, 1).text()
                for data in self.filtered_data:
                    if data[0] == billing_code:
                        selected_rows.append(data)
                        break

        if not selected_rows:
            QtWidgets.QMessageBox.warning(self, "No Selection", "Please select at least one billing to issue.")
            return

        reply = QtWidgets.QMessageBox.question(
            self,
            "Confirm Issue",
            f"Are you sure you want to issue {len(selected_rows)} bill(s)?",
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
        )
        if reply == QtWidgets.QMessageBox.Yes:
            for billing in selected_rows:
                # Call existing confirm_issue_bill without editing it
                self.parent_window.confirm_issue_bill(self, billing)

            self.accept()

class QuickSetPaidDialog(QtWidgets.QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Quick Set Paid")
        self.setMinimumSize(800, 500)
        self.billing_data = []
        self.filtered_data = []
        self.parent_window = parent

        self.setup_ui()
        self.load_billing_data()

    def setup_ui(self):
        layout = QtWidgets.QVBoxLayout(self)

        # Search Bar
        search_layout = QtWidgets.QHBoxLayout()
        self.search_input = QtWidgets.QLineEdit()
        self.search_input.setPlaceholderText("Search by Billing Code or Client Name...")
        self.search_input.textChanged.connect(self.filter_table)
        search_layout.addWidget(QtWidgets.QLabel("Search:"))
        search_layout.addWidget(self.search_input)

        layout.addLayout(search_layout)

        # Table Setup
        self.table = QtWidgets.QTableWidget()
        self.table.setColumnCount(3)
        self.table.setHorizontalHeaderLabels(["Select", "Billing Code", "Client Name"])
        self.table.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.Stretch)
        self.table.setSortingEnabled(True)
        layout.addWidget(self.table)

        # Button Row
        btn_layout = QtWidgets.QHBoxLayout()
        self.paid_btn = QtWidgets.QPushButton("✅ Set Selected as Paid")
        self.paid_btn.clicked.connect(self.on_paid_clicked)
        btn_layout.addStretch()
        btn_layout.addWidget(self.paid_btn)

        layout.addLayout(btn_layout)

    def load_billing_data(self):
        backend = adminPageBack()
        raw_data = backend.fetch_billing_pending_payment()  # We'll define this
        self.billing_data = raw_data
        self.filtered_data = raw_data
        self.update_table()

    def filter_table(self):
        query = self.search_input.text().strip().lower()
        if not query:
            self.filtered_data = self.billing_data
        else:
            self.filtered_data = [
                row for row in self.billing_data
                if query in row[0].lower() or query in row[2].lower()
            ]
        self.update_table()

    def update_table(self):
        self.table.setRowCount(len(self.filtered_data))
        for i, (billing_code, issued_date, client_name) in enumerate(self.filtered_data):
            checkbox = QtWidgets.QTableWidgetItem()
            checkbox.setFlags(QtCore.Qt.ItemIsUserCheckable | QtCore.Qt.ItemIsEnabled)
            checkbox.setCheckState(QtCore.Qt.Unchecked)

            self.table.setItem(i, 0, checkbox)
            self.table.setItem(i, 1, QtWidgets.QTableWidgetItem(billing_code))
            self.table.setItem(i, 2, QtWidgets.QTableWidgetItem(client_name))

    def on_paid_clicked(self):
        selected_rows = []
        for i in range(self.table.rowCount()):
            item = self.table.item(i, 0)
            if item.checkState() == QtCore.Qt.Checked:
                billing_code = self.table.item(i, 1).text()
                for data in self.filtered_data:
                    if data[0] == billing_code:
                        selected_rows.append(data)
                        break

        if not selected_rows:
            QtWidgets.QMessageBox.warning(self, "No Selection", "Please select at least one billing to mark as paid.")
            return

        reply = QtWidgets.QMessageBox.question(
            self,
            "Confirm Payment",
            f"Mark {len(selected_rows)} bill(s) as PAID?",
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
        )
        if reply == QtWidgets.QMessageBox.Yes:
            for billing in selected_rows:
                # Call existing mark_billing_as_paid without editing it
                self.parent_window.mark_billing_as_paid(billing)

            self.accept()


class EmployeeBillingPage(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super().__init__()
        self.parent = parent
        
        # Store the username from parent
        self.username = parent if isinstance(parent, str) else "System"
        
        # Initialize pagination variables before setup_ui
        self.rows_per_page = 10
        self.current_page = 1
        self.total_pages = 1
        self.all_data = []  # Store all data for pagination
        self.filtered_data = []  # Store filtered data for pagination
        
        # Now call setup_ui after initializing variables
        self.selected_status = "ALL"
        self.setup_ui()

    def create_scrollable_cell(self, row, column, text):
        scrollable_widget = ScrollableTextWidget(text)
        self.billing_table.setCellWidget(row, column, scrollable_widget)

    def create_action_cell(self, row, billing_data):
        action_widget = QtWidgets.QWidget()
        action_layout = QtWidgets.QHBoxLayout(action_widget)
        action_layout.setContentsMargins(5, 5, 5, 5)
        action_layout.setAlignment(QtCore.Qt.AlignCenter)

        status = billing_data[7]

        # Print button for TO BE PRINTED
        if status == "TO BE PRINTED":
            print_btn = QtWidgets.QPushButton("🖨 Print")
            print_btn.setStyleSheet("""
                QPushButton {
                    background-color: #4CAF50;
                    color: white;
                    padding: 5px 10px;
                    border-radius: 4px;
                }
                QPushButton:hover {
                    background-color: #45A049;
                }
            """)
            print_btn.clicked.connect(lambda _, data=billing_data: self.print_bill(data))
            action_layout.addWidget(print_btn)

        # View button for PRINTED
        if status in ["PRINTED", "VOID", "PENDING PAYMENT"]:
            view_btn = QtWidgets.QPushButton("📄 View")
            view_btn.setStyleSheet("""
                QPushButton {
                    background-color: #2196F3;
                    color: white;
                    padding: 5px 10px;
                    border-radius: 4px;
                }
                QPushButton:hover {
                    background-color: #1976D2;
                }
            """)
            view_btn.clicked.connect(lambda _, data=billing_data: self.view_bill(data))
            action_layout.addWidget(view_btn)

        # Edit button for TO BE PRINTED and PRINTED
        if status in ["TO BE PRINTED", "PRINTED"]:
            edit_btn = QtWidgets.QPushButton("✏️ Edit")
            edit_btn.setStyleSheet("""
                QPushButton {
                    background-color: #FFB74D;
                    color: white;
                    padding: 5px 10px;
                    border-radius: 4px;
                }
                QPushButton:hover {
                    background-color: #FFA726;
                }
            """)
            edit_btn.clicked.connect(lambda _, data=billing_data: self.edit_billing(data))
            action_layout.addWidget(edit_btn)

        # Reissue button for PENDING PAYMENT
        if status == "PENDING PAYMENT":
            paid_btn = QtWidgets.QPushButton("💰 Paid")
            paid_btn.setStyleSheet("""
                QPushButton {
                    background-color: #4CAF50;
                    color: white;
                    padding: 5px 10px;
                    border-radius: 4px;
                }
                QPushButton:hover {
                    background-color: #388E3C;
                }
            """)
            paid_btn.clicked.connect(lambda _, data=billing_data: self.mark_billing_as_paid(data))
            action_layout.addWidget(paid_btn)

            reissue_btn = QtWidgets.QPushButton("♻️ Reissue")
            reissue_btn.setStyleSheet("""
                QPushButton {
                    background-color: #FF9800;
                    color: white;
                    padding: 5px 10px;
                    border-radius: 4px;
                }
                QPushButton:hover {
                    background-color: #F57C00;
                }
            """)
            reissue_btn.clicked.connect(lambda _, data=billing_data: self.reissue_billing(data))
            action_layout.addWidget(reissue_btn)

            void_btn = QtWidgets.QPushButton("🚫 Void")
            void_btn.setStyleSheet("""
                QPushButton {
                    background-color: #BDBDBD;
                    color: white;
                    padding: 5px 10px;
                    border-radius: 4px;
                }
                QPushButton:hover {
                    background-color: #9E9E9E;
                }
            """)
            void_btn.clicked.connect(lambda _, data=billing_data: self.void_billing(data))
            action_layout.addWidget(void_btn)


        # Delete button for TO BE PRINTED
        if status == "TO BE PRINTED":
            delete_btn = QtWidgets.QPushButton("❌ Delete")
            delete_btn.setStyleSheet("""
                QPushButton {
                    background-color: #E57373;
                    color: white;
                    padding: 5px 10px;
                    border-radius: 4px;
                }
                QPushButton:hover {
                    background-color: #EF5350;
                }
            """)
            delete_btn.clicked.connect(lambda _, data=billing_data: self.delete_billing(data))
            action_layout.addWidget(delete_btn)

        # Void button for PRINTED
        if status == "PRINTED":
            void_btn = QtWidgets.QPushButton("🚫 Void")
            void_btn.setStyleSheet("""
                QPushButton {
                    background-color: #BDBDBD;
                    color: white;
                    padding: 5px 10px;
                    border-radius: 4px;
                }
                QPushButton:hover {
                    background-color: #9E9E9E;
                }
            """)
            void_btn.clicked.connect(lambda _, data=billing_data: self.void_billing(data))
            action_layout.addWidget(void_btn)

        self.billing_table.setCellWidget(row, 8, action_widget)

    def mark_billing_as_paid(self, billing_data):
        reply = QtWidgets.QMessageBox.question(
            self, "Confirm Payment",
            f"Mark billing {billing_data[0]} and its transaction as PAID?",
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
        )
        if reply == QtWidgets.QMessageBox.Yes:
            try:
                backend = adminPageBack()
                billing_code = billing_data[0]
                billing_id = backend.get_billing_id(billing_code)

                # Update billing status
                backend.update_billing_status(billing_id, "PAID")

                # Get the transaction ID linked to the billing
                transaction_id = backend.get_transaction_id_by_billing_id(billing_id)
                if transaction_id:
                    backend.mark_transaction_paid(transaction_id, date.today())


                QtWidgets.QMessageBox.information(self, "Success", f"Billing {billing_code} and its transaction marked as PAID.")
                self.populate_table(backend.fetch_billing())
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, "Error", f"Failed to update payment status: {e}")


    # Implement edit, delete, and void methods
    def edit_billing(self, billing_data):
        backend = adminPageBack()
        billing_code = billing_data[0]
        billing_id = backend.get_billing_id(billing_code)
        billing_details = backend.get_billing_by_id(billing_id)

        if not billing_details:
            QtWidgets.QMessageBox.warning(self, "Edit Billing", f"No billing data found for {billing_code}")
            return

        self.show_edit_billing(existing_data=billing_details)

    
    def show_edit_billing(self, existing_data):
        if not existing_data or len(existing_data) == 0:
            QtWidgets.QMessageBox.warning(self, "Edit Billing", "No billing data available for editing.")
            return

        dialog = QtWidgets.QDialog(self)
        dialog.setWindowTitle("Edit Billing")
        dialog.setFixedSize(1000, 700)
        dialog.setStyleSheet("background-color: #C9EBCB;")

        layout = QtWidgets.QVBoxLayout(dialog)
        layout.setContentsMargins(30, 5, 30, 5)

        title = QtWidgets.QLabel("EDIT BILLING INFORMATION")
        title.setStyleSheet("font-size: 20px; padding: 10px;")
        title.setAlignment(QtCore.Qt.AlignCenter)
        layout.addWidget(title)

        input_style = """
            QLineEdit, QDateEdit, QComboBox {
                font-family: 'Arial';
                font-size: 14px;
                padding: 8px;
                border: 1px solid #bdc3c7;
                border-radius: 4px;
                background-color: #ffffff;
            }
        """
        readonly_style = """
            QLineEdit {
                font-family: 'Arial';
                font-size: 14px;
                padding: 8px;
                border: 1px solid #bdc3c7;
                border-radius: 4px;
                background-color: #e0e0e0;
                color: #555555;
            }
        """

        billing_id, billing_due_date, billing_total_amount, cubic_meter_val, reading_id, client_id, categ_id, billing_date, billing_code, billing_status, issued_date, amount_val, subscribe_capital_val, late_payment_val, penalty_val, total_charge_val = existing_data

        backend = adminPageBack()
        client_info = backend.fetch_client_by_id(client_id)[0]
        client_name_val = f"{client_info[2]}, {client_info[1]}"
        client_location_val = client_info[4]
        meter_number_val = client_info[5]
        client_categ_id = client_info[7]
        self.rate_blocks = backend.fetch_rate_blocks_by_categ(client_categ_id)


        reading_info = backend.get_reading_info_by_id(reading_id)
        reading_date_val = reading_info[1]
        reading_current_val = reading_info[3]  # current reading

        client_location = QtWidgets.QLineEdit(client_location_val)
        client_location.setStyleSheet(readonly_style)
        client_location.setReadOnly(True)

        client_name = QtWidgets.QLineEdit(client_name_val)
        client_name.setStyleSheet(readonly_style)
        client_name.setReadOnly(True)

        meter_number = QtWidgets.QLineEdit(str(meter_number_val))
        meter_number.setStyleSheet(readonly_style)
        meter_number.setReadOnly(True)

        reading_date = QtWidgets.QDateEdit()
        reading_date.setDate(reading_date_val)
        reading_date.setCalendarPopup(True)
        reading_date.setStyleSheet(input_style)
        reading_date.setReadOnly(False)  # optional, since QDateEdit allows changes

        previous_reading = QtWidgets.QLineEdit(str(reading_info[2]))  # reading_prev
        previous_reading.setStyleSheet(readonly_style)
        previous_reading.setReadOnly(True)


        current_reading = QtWidgets.QLineEdit(str(reading_current_val))
        current_reading.setStyleSheet(input_style)



        subscribe_capital = QtWidgets.QLineEdit(str(subscribe_capital_val))
        subscribe_capital.setStyleSheet(input_style)

        late_payment = QtWidgets.QLineEdit(str(late_payment_val))
        late_payment.setStyleSheet(input_style)

        penalty = QtWidgets.QLineEdit(str(penalty_val))
        penalty.setStyleSheet(input_style)

        total_charge = QtWidgets.QLineEdit(str(total_charge_val))
        total_charge.setStyleSheet(readonly_style)
        total_charge.setReadOnly(True)

        amount = QtWidgets.QLineEdit(str(amount_val))
        amount.setStyleSheet(readonly_style)
        amount.setReadOnly(True)

        total_bill = QtWidgets.QLineEdit(str(billing_total_amount))
        total_bill.setStyleSheet(readonly_style)
        total_bill.setReadOnly(True)

        cubic_meter = QtWidgets.QLineEdit()
        cubic_meter.setStyleSheet(readonly_style)
        cubic_meter.setReadOnly(True)
        cubic_meter.setText(str(cubic_meter_val))



        billing_due = QtWidgets.QDateEdit()
        billing_due.setDate(billing_due_date)
        billing_due.setCalendarPopup(True)
        billing_due.setStyleSheet(input_style)

        form_layout = QtWidgets.QGridLayout()
        form_layout.setHorizontalSpacing(30)
        form_layout.setVerticalSpacing(8)
        form_layout.setColumnStretch(0, 1)  # Left column
        form_layout.setColumnStretch(1, 1)

        # LEFT COLUMN
        form_layout.addWidget(QtWidgets.QLabel("CLIENT NAME:"), 0, 0)
        form_layout.addWidget(client_name, 1, 0)

        form_layout.addWidget(QtWidgets.QLabel("READING DATE:"), 4, 0)
        form_layout.addWidget(reading_date, 5, 0)

        form_layout.addWidget(QtWidgets.QLabel("PREVIOUS READING:"), 6, 0)
        form_layout.addWidget(previous_reading, 7, 0)

        form_layout.addWidget(QtWidgets.QLabel("CURRENT READING:"), 8, 0)
        form_layout.addWidget(current_reading, 9, 0)

        form_layout.addWidget(QtWidgets.QLabel("CUBIC METER CONSUMED:"), 10, 0)
        form_layout.addWidget(cubic_meter, 11, 0)

        form_layout.addWidget(QtWidgets.QLabel("AMOUNT:"), 12, 0)
        form_layout.addWidget(amount, 13, 0)

        form_layout.addWidget(QtWidgets.QLabel("DUE DATE:"), 14, 0)
        form_layout.addWidget(billing_due, 15, 0)


        # RIGHT COLUMN
        additional_label = QtWidgets.QLabel("ADDITIONAL CHARGE")
        font = QtGui.QFont()
        font.setBold(True)
        font.setPointSize(11)
        additional_label.setFont(font)
        additional_label.setAlignment(QtCore.Qt.AlignCenter)
        form_layout.addWidget(additional_label, 0, 1, 1, 1)

        form_layout.addWidget(QtWidgets.QLabel("SUBSCRIBE CAPITAL:"), 1, 1)
        form_layout.addWidget(subscribe_capital, 2, 1)

        form_layout.addWidget(QtWidgets.QLabel("LATE PAYMENT:"), 3, 1)
        form_layout.addWidget(late_payment, 4, 1)

        form_layout.addWidget(QtWidgets.QLabel("PENALTY:"), 5, 1)
        form_layout.addWidget(penalty, 6, 1)

        form_layout.addWidget(QtWidgets.QLabel("TOTAL CHARGE:"), 7, 1)
        form_layout.addWidget(total_charge, 8, 1)

        form_layout.addWidget(QtWidgets.QLabel("TOTAL BILL:"), 9, 1)
        form_layout.addWidget(total_bill, 10, 1)



        layout.addLayout(form_layout)

        button_container = QtWidgets.QWidget()
        button_layout = QtWidgets.QHBoxLayout(button_container)
        button_layout.setAlignment(QtCore.Qt.AlignRight)

        cancel_btn = QtWidgets.QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        cancel_btn.setStyleSheet("padding:8px; background-color: #95a5a6; color:white; border-radius:4px;")

        save_btn = QtWidgets.QPushButton("Save Changes")
        save_btn.setStyleSheet("padding:8px; background-color: #27ae60; color:white; border-radius:4px;")

        button_layout.addWidget(cancel_btn)
        button_layout.addWidget(save_btn)
        layout.addWidget(button_container)

        # Real-time update logic for total_charge and total_bill
        def update_totals():
            try:
                sub_cap = float(subscribe_capital.text()) if subscribe_capital.text() else 0
                late_pay = float(late_payment.text()) if late_payment.text() else 0
                pen = float(penalty.text()) if penalty.text() else 0
                charge_total = sub_cap + late_pay + pen
                total_charge.setText(f"{charge_total:.2f}")

                amt = float(amount.text()) if amount.text() else 0
                total_bill_val = amt + charge_total
                total_bill.setText(f"{total_bill_val:.2f}")
            except ValueError:
                total_charge.setText("0.00")
                total_bill.setText("0.00")

        #realtime update cubic meter consumed
        def update_cubic_meter():
            try:
                prev = float(previous_reading.text())
                curr = float(current_reading.text())
                consumed = max(0, curr - prev)
                cubic_meter.setText(f"{consumed:.2f}")
            except ValueError:
                cubic_meter.setText("0.00")

        def on_current_reading_changed():
            try:
                prev = float(previous_reading.text())
                curr = float(current_reading.text())

                if curr < prev:
                    cubic_meter.setText("0")
                    amount.setText("0.00")
                    return

                consumed = curr - prev
                cubic_meter.setText(str(consumed))

                total_amount = 0

                for block in self.rate_blocks:
                    is_minimum = block[1]
                    min_c = block[2]
                    max_c = block[3] if block[3] is not None else float('inf')
                    rate = block[4]

                    if is_minimum:
                        total_amount += rate
                    elif not is_minimum and consumed >= min_c:
                        applied_volume = max(0, min(consumed, max_c) - min_c + 1)
                        total_amount += applied_volume * rate

                amount.setText(f"{total_amount:.2f}")
                update_total_bill()

            except ValueError:
                cubic_meter.setText("0")
                amount.setText("0.00")

        def update_total_bill():
            try:
                sub_cap = float(subscribe_capital.text()) if subscribe_capital.text() else 0
                late_pay = float(late_payment.text()) if late_payment.text() else 0
                pen = float(penalty.text()) if penalty.text() else 0
                charge_total = sub_cap + late_pay + pen
                total_charge.setText(f"{charge_total:.2f}")

                amt = float(amount.text()) if amount.text() else 0
                total_bill_val = amt + charge_total
                total_bill.setText(f"{total_bill_val:.2f}")
            except ValueError:
                total_charge.setText("0.00")
                total_bill.setText("0.00")




        subscribe_capital.textChanged.connect(update_totals)
        late_payment.textChanged.connect(update_totals)
        penalty.textChanged.connect(update_totals)
        current_reading.textChanged.connect(update_cubic_meter)
        current_reading.textChanged.connect(on_current_reading_changed)



        def save_edited_billing():
            backend = adminPageBack()

            try:
                updated_total = float(total_bill.text())
                updated_due = billing_due.date().toPyDate()
                updated_sub_capital = float(subscribe_capital.text())
                updated_late_payment = float(late_payment.text())
                updated_penalty = float(penalty.text())
                updated_total_charge = float(total_charge.text())

                # ➕ New: Update Reading
                updated_reading_date = reading_date.date().toPyDate()
                updated_current_reading = float(current_reading.text())
                meter_id = backend.get_meter_id_by_reading_id(reading_id)

                # Update reading table (only current reading and date)
                backend.update_reading(reading_id, updated_reading_date, updated_current_reading)

                # ➕ New: Update meter with current reading and date
                backend.update_meter_latest_reading(updated_current_reading, updated_reading_date, meter_id)

                # ✅ Update billing table
                backend.edit_billing(
                    billing_id,
                    updated_total,
                    updated_due,
                    updated_sub_capital,
                    updated_late_payment,
                    updated_penalty,
                    updated_total_charge,
                    float(amount.text()),                  # billing_amount
                    float(cubic_meter.text()),             # billing_consumption
                    reading_date.date().toPyDate()         # billing_date
                )


                QtWidgets.QMessageBox.information(dialog, "Updated", "Billing and corresponding data updated successfully.")
                dialog.accept()
                self.populate_table(backend.fetch_billing())

            except Exception as e:
                QtWidgets.QMessageBox.critical(dialog, "Error", f"Failed to update billing: {e}")

        save_btn.clicked.connect(save_edited_billing)
        dialog.exec_()

    def delete_billing(self, billing_data):
        reply = QtWidgets.QMessageBox.question(
            self, "Confirm Deletion",
            f"Are you sure you want to delete billing {billing_data[0]}?",
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
        )
        if reply == QtWidgets.QMessageBox.Yes:
            try:
                adminPageBack().delete_billing(billing_data[0])
                QtWidgets.QMessageBox.information(self, "Deleted",
                                                  f"Billing {billing_data[0]} and its related reading were deleted.")
                self.populate_table(adminPageBack().fetch_billing())
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, "Error", f"Failed to delete billing: {e}")


    def void_billing(self, billing_data):
        reply = QtWidgets.QMessageBox.question(
            self, "Confirm Void",
            f"Are you sure you want to void billing {billing_data[0]} and its associated reading and transaction?",
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
        )
        if reply == QtWidgets.QMessageBox.Yes:
            try:
                backend = adminPageBack()
                billing_code = billing_data[0]
                billing_id = backend.get_billing_id(billing_code)

                # Mark billing as VOID
                backend.update_billing_status(billing_id, "VOID")

                # Get reading_id and transaction_id
                reading_id = backend.get_reading_id_by_billing_id(billing_id)
                trans_id = backend.get_transaction_id_by_billing_id(billing_id)

                # Mark transaction as voided
                if trans_id:
                    backend.update_transaction_status(trans_id, "VOID")

                # Mark reading as voided and update meter with previous reading
                if reading_id:
                    reading_info = backend.get_reading_info_by_id(reading_id)
                    if reading_info:
                        reading_prev = reading_info[2]      # reading_prev
                        reading_date = reading_info[1]     # reading_date
                        meter_id = reading_info[4]         # meter_id

                        backend.void_reading(reading_id)
                        # Find matching non-voided reading with current = voided reading_prev
                        matching_reading = backend.get_reading_by_current_and_meter(reading_prev, meter_id)

                        print("matching_reading result:", matching_reading)
                        print("type:", type(matching_reading))

                        if matching_reading:
                            # reading_current matched, use its date
                            updated_reading_date = matching_reading[1]
                        else:
                            # fallback: use current reading_date from voided reading
                            updated_reading_date = reading_date

                        # Now update the meter with adjusted last reading and date
                        backend.update_meter_latest_reading(reading_prev, updated_reading_date, meter_id)


                QtWidgets.QMessageBox.information(self, "Voided", f"Billing {billing_code} and its related data voided successfully.")
                self.populate_table(backend.fetch_billing())
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, "Error", f"Failed to void billing and related data: {e}")







    class LoadingDialog(QtWidgets.QDialog):
        def __init__(self, message="Processing..."):
            super().__init__()
            self.setWindowFlags(QtCore.Qt.Dialog | QtCore.Qt.FramelessWindowHint)
            self.setModal(True)
            self.setAttribute(QtCore.Qt.WA_TranslucentBackground)

            # Main container
            container = QtWidgets.QFrame()
            container.setStyleSheet("""
                QFrame {
                    background-color: #ffffff;
                    border-radius: 16px;
                    border: none;
                    box-shadow: 0px 4px 12px rgba(0, 0, 0, 0.1);
                }
            """)

            layout = QtWidgets.QVBoxLayout(self)
            layout.setContentsMargins(15, 15, 15, 15)
            layout.addWidget(container)

            inner_layout = QtWidgets.QVBoxLayout(container)
            inner_layout.setSpacing(15)
            inner_layout.setContentsMargins(20, 20, 20, 20)
            

            # Spinner icon or fallback
            self.spinner = QtWidgets.QLabel()
            self.spinner.setFixedSize(64, 64)
            self.spinner.setAlignment(QtCore.Qt.AlignCenter)

            # Absolute path to GIF
            spinner_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "gifs", "Spinner.gif"))
            print("Spinner path:", spinner_path)

            self.movie = QtGui.QMovie(spinner_path)
            if not self.movie.isValid():
                self.spinner.setText("⏳ Loading...")
            else:
                self.movie.setScaledSize(QtCore.QSize(64, 64))
                self.spinner.setMovie(self.movie)
                self.movie.start()


            # Message label
            self.label = QtWidgets.QLabel(message)
            self.label.setAlignment(QtCore.Qt.AlignCenter)
            self.label.setStyleSheet("font-size: 16px; color: #333333;")

            # Progress bar
            self.progress = QtWidgets.QProgressBar()
            self.progress.setRange(0, 0)
            self.progress.setTextVisible(False)
            self.progress.setStyleSheet("""
                QProgressBar {
                    background-color: #e0e0e0;
                    border: none;
                    border-radius: 6px;
                    height: 14px;
                }
                QProgressBar::chunk {
                    background-color: #4CAF50;
                    border-radius: 6px;
                }
            """)
            inner_layout.addStretch()
            inner_layout.addWidget(self.spinner, alignment=QtCore.Qt.AlignCenter)
            inner_layout.addWidget(self.label, alignment=QtCore.Qt.AlignCenter)
            inner_layout.addWidget(self.progress)
            inner_layout.addStretch()


            self.setFixedSize(320, 240)

        def set_message(self, text):
            self.label.setText(text)

        def set_progress(self, value, max_value=100):
            self.progress.setRange(0, max_value)
            self.progress.setValue(value)
            QtWidgets.QApplication.processEvents()


    def generate_pdf_path(self, billing_data):
        """Generate the file path for the PDF based on billing data."""
        # Generate the billing code (or any unique identifier for the bill)
        billing_code = billing_data[0]  # assuming the first element is the billing code
        
        # Get the base directory for saving the PDF
        base_dir = os.path.abspath(os.path.join(os.path.dirname(sys.argv[0]), ".."))
        
        # Create the path for the PDF file
        pdf_path = os.path.join(base_dir, "temp_bills", f"bill_{billing_code}.pdf")
        
        # Ensure the directory exists, if not, create it
        os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
        print(f"PDF Path: {pdf_path}")
        return pdf_path


   


    def fill_word_template(self, data, template_path="bill_template.docx", output_path="generated_bill.docx"):
        doc = Document(template_path)

        # Add checkmarks only for the correct category
        for i in range(1, 5):
            data[f"check{i}"] = "✓" if data.get("category_code") == i else ""

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))

        doc.save(output_path)
        print(f"✅ Word file generated at: {output_path}")

    def convert_to_pdf(self,docx_path="generated_bill.docx", pdf_path="generated_bill.pdf"):
        if os.path.exists(pdf_path):
            os.remove(pdf_path)  # Delete old version
        convert(docx_path, pdf_path)
        print(f"✅ PDF saved to {pdf_path}")

    class ViewBill(QWidget):
        def __init__(self, pdf_path, batch_mode=False):
            super().__init__()
            self.pdf_path = pdf_path
            self.batch_mode = batch_mode

            self.setWindowTitle("View Bill")
            self.setMinimumSize(900, 700)

            # Scrollable layout
            scroll = QScrollArea(self)
            layout = QVBoxLayout(self)
            layout.addWidget(scroll)

            container = QWidget()
            scroll.setWidget(container)
            scroll.setWidgetResizable(True)

            container_layout = QVBoxLayout(container)

            try:
                # Load PDF pages as images using PyMuPDF (fitz)
                try:
                    doc = fitz.open(pdf_path)
                except Exception as e:
                    error_label = QLabel(f"❌ Failed to open PDF: {e}")
                    container_layout.addWidget(error_label)
                    return
                for i, page in enumerate(doc):
                    try:
                        # Render page to image
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        image_format = QImage.Format_RGBA8888 if pix.alpha else QImage.Format_RGB888
                        qimage = QImage(pix.samples, pix.width, pix.height, pix.stride, image_format)
                        pixmap = QPixmap.fromImage(qimage)

                        # Add label for batch mode
                        if self.batch_mode:
                            page_label = QLabel(f"📄 - BILL {i + 1}")
                            page_label.setAlignment(QtCore.Qt.AlignCenter)
                            page_label.setStyleSheet("color: red; font-size: 14px; font-weight: bold; padding: 10px;")
                            container_layout.addWidget(page_label)

                        # Show rendered image
                        image_label = QLabel()
                        image_label.setPixmap(pixmap)
                        image_label.setAlignment(QtCore.Qt.AlignCenter)
                        container_layout.addWidget(image_label)

                        # Reprint button (batch mode only)
                        if self.batch_mode:
                            reprint_button = QPushButton(f"🖨 Reprint BILL {i + 1}")
                            reprint_button.setStyleSheet("""
                                QPushButton {
                                    background-color: #4CAF50;
                                    color: white;
                                    font-weight: bold;
                                    padding: 6px 12px;
                                    border-radius: 4px;
                                }
                                QPushButton:hover {
                                    background-color: #45A049;
                                }
                            """)
                            reprint_button.clicked.connect(lambda _, idx=i: self.print_single_page(idx))
                            container_layout.addWidget(reprint_button)

                    except Exception as page_error:
                        error_label = QLabel(f"❌ Failed to load page {i + 1}: {page_error}")
                        container_layout.addWidget(error_label)

                    # Spacing between each bill
                    spacer = QLabel()
                    spacer.setFixedHeight(40)
                    container_layout.addWidget(spacer)

            except Exception as e:
                error_label = QLabel(f"❌ Failed to load preview:\n{str(e)}")
                error_label.setAlignment(QtCore.Qt.AlignCenter)
                error_label.setStyleSheet("color: red; font-weight: bold; padding: 20px;")
                container_layout.addWidget(error_label)

    def view_bill(self, billing_data):

        self.last_viewed_bill_data = billing_data

        try:
            # Set up base paths
            base_dir = os.path.abspath(os.path.join(os.path.dirname(sys.argv[0]), ".."))
            template_path = os.path.join(base_dir, "templates", "preview_template.docx")
            temp_folder = os.path.join(base_dir, "temp_view_bill")
            os.makedirs(temp_folder, exist_ok=True)

            # 🔄 Clear previous files
            for f in os.listdir(temp_folder):
                file_path = os.path.join(temp_folder, f)
                if os.path.isfile(file_path):
                    os.unlink(file_path)

            # ✅ Show loading dialog
            self.loading = self.LoadingDialog("Loading Bill...")
            self.loading.set_progress(0, 1)
            self.loading.show()

            # ✅ Background worker (reusing the same logic)
            worker = BillWorker(
                [billing_data],
                template_path,
                temp_folder,
                self.fill_word_template,
                self.convert_to_pdf
            )
            worker.signals.progress.connect(lambda i, total: (
                self.loading.set_message(f"Loading..."),
                self.loading.set_progress(i, total)
            ))
            worker.signals.finished.connect(self.on_solo_view_generated)
            worker.signals.error.connect(self.on_generation_failed)

            QThreadPool.globalInstance().start(worker)

        except Exception as e:
            if hasattr(self, "loading"):
                self.loading.close()
            QtWidgets.QMessageBox.warning(self, "Bill Preview Failed", str(e))

    def view_all_on_page(self):

        # Base path setup
        base_dir = os.path.abspath(os.path.join(os.path.dirname(sys.argv[0]), ".."))
        template_path = os.path.join(base_dir, "templates", "preview_template.docx")

        if not os.path.exists(template_path):
            QtWidgets.QMessageBox.warning(self, "Template Missing", f"Template not found at:\n{template_path}")
            return

        # Prepare temp folder
        temp_folder = os.path.join(base_dir, "temp_view_bill")
        os.makedirs(temp_folder, exist_ok=True)

        # Clear old files
        for f in os.listdir(temp_folder):
            file_path = os.path.join(temp_folder, f)
            if os.path.isfile(file_path):
                os.unlink(file_path)


        # Get current page data
        start_idx = (self.current_page - 1) * self.rows_per_page
        end_idx = min(start_idx + self.rows_per_page, len(self.filtered_data))
        page_data = self.filtered_data[start_idx:end_idx]

        # Show animated loading dialog
        self.loading = self.LoadingDialog("Loading Bills...")
        self.loading.set_progress(0, len(page_data))
        self.loading.show()

        # Start background generation
        worker = BillWorker(page_data, template_path, temp_folder, self.fill_word_template, self.convert_to_pdf)
        worker.signals.progress.connect(lambda i, total: (
            self.loading.set_message(f"Generating BILL {i} of {total}"),
            self.loading.set_progress(i, total)
        ))
        worker.signals.finished.connect(self.on_generation_finished)
        worker.signals.error.connect(self.on_generation_failed)

        QThreadPool.globalInstance().start(worker)

    def print_bill(self, billing_data):

        self.last_printed_bill_data = billing_data
        
        try:
            # Set up base paths
            base_dir = os.path.abspath(os.path.join(os.path.dirname(sys.argv[0]), ".."))
            template_path = os.path.join(base_dir, "templates", "bill_template.docx")
            temp_folder = os.path.join(base_dir, "temp_single_bill")
            os.makedirs(temp_folder, exist_ok=True)

            # 🔄 Clear previous files
            for f in os.listdir(temp_folder):
                file_path = os.path.join(temp_folder, f)
                if os.path.isfile(file_path):
                    os.unlink(file_path)

            # ✅ Show loading dialog
            self.loading = self.LoadingDialog("Generating bill...")
            self.loading.set_progress(0, 1)
            self.loading.show()

            # ✅ Background worker (reusing the same logic)
            worker = BillWorker(
                [billing_data],
                template_path,
                temp_folder,
                self.fill_word_template,
                self.convert_to_pdf
            )
            worker.signals.progress.connect(lambda i, total: (
                self.loading.set_message(f"Generating bill..."),
                self.loading.set_progress(i, total)
            ))
            worker.signals.finished.connect(self.on_solo_generation_finished)
            worker.signals.error.connect(self.on_generation_failed)
            
            QThreadPool.globalInstance().start(worker)

        except Exception as e:
            if hasattr(self, "loading"):
                self.loading.close()
            QtWidgets.QMessageBox.warning(self, "Bill Generation Failed", str(e))

    def on_solo_view_generated(self, merged_pdf_path):
        self.loading.close()
        self.preview_window = self.ViewBill(merged_pdf_path)

        billing_data = self.last_viewed_bill_data
        billing_status = billing_data[7]  # Status column

        if billing_status not in ["VOID", "PENDING PAYMENT"]:
            # ✅ Reprint Button
            print_btn = QtWidgets.QPushButton("🖨 Reprint")
            print_btn.setStyleSheet("""
                QPushButton {
                    background-color: #4CAF50;
                    color: white;
                    padding: 6px 12px;
                    border-radius: 4px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #388E3C;
                }
            """)
            print_btn.clicked.connect(lambda: self.preview_window.print_pdf())
            self.preview_window.layout().addWidget(print_btn)

            # ✅ Issue Button
            issue_btn = QtWidgets.QPushButton("✅ Issue")
            issue_btn.setStyleSheet("""
                QPushButton {
                    background-color: #43A047;
                    color: white;
                    padding: 6px 12px;
                    border-radius: 4px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #2E7D32;
                }
            """)
            issue_btn.clicked.connect(lambda: self.confirm_issue_bill(self.preview_window, billing_data))
            self.preview_window.layout().addWidget(issue_btn)

            # ✅ Void Button
            void_btn = QtWidgets.QPushButton("🚫 Void")
            void_btn.setStyleSheet("""
                QPushButton {
                    background-color: #E53935;
                    color: white;
                    padding: 6px 12px;
                    border-radius: 4px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #C62828;
                }
            """)
            void_btn.clicked.connect(lambda: self.confirm_void_bill(self.preview_window, billing_data))
            self.preview_window.layout().addWidget(void_btn)

        self.preview_window.show()



    def on_view_generated(self, merged_pdf_path):
        self.loading.close()
        self.preview_window = self.ViewBill(merged_pdf_path, batch_mode=True)

        # Mark All as Printed button
        mark_all_btn = QtWidgets.QPushButton("✅ Mark All as Printed")
        mark_all_btn.setStyleSheet("""
            QPushButton {
                background-color: #388E3C;
                color: white;
                padding: 6px 12px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2E7D32;
            }
        """)
        start_idx = (self.current_page - 1) * self.rows_per_page
        end_idx = min(start_idx + self.rows_per_page, len(self.filtered_data))
        page_data = self.filtered_data[start_idx:end_idx]
        mark_all_btn.clicked.connect(lambda: self.mark_as_printed(self.preview_window, page_data))

        self.preview_window.layout().addWidget(mark_all_btn)
        self.preview_window.show()

    
    def confirm_issue_bill(self, preview_window, billing_data):
        reply = QtWidgets.QMessageBox.question(
            preview_window,
            "Confirm Issue",
            f"Are you sure you want to issue billing {billing_data[0]} and mark it as PENDING PAYMENT?",
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
        )
        if reply == QtWidgets.QMessageBox.Yes:
            try:
                backend = adminPageBack()
                billing_code = billing_data[0]
                billing_id = backend.get_billing_id(billing_code)
                bill_details = backend.get_billing_by_id(billing_id)

                from datetime import date
                today = date.today()

                # 🔄 Update status and issued date
                backend.update_billing_status(billing_id, "PENDING PAYMENT")
                backend.update_billing_issued_date(billing_id, today)

                # ✅ Create Transaction
                trans_status = "PENDING"
                trans_payment_date = None
                trans_total_amount = bill_details[2]  # billing_total
                client_id = bill_details[5]
                reading_id = bill_details[4]  # assuming this is the reading_id

                backend.add_transaction(
                    billing_id, trans_status, trans_payment_date, trans_total_amount,
                    client_id, reading_id
                )

                QtWidgets.QMessageBox.information(preview_window, "Issued", f"Billing {billing_code} marked as PENDING PAYMENT and transaction created.")
                preview_window.close()
                self.populate_table(backend.fetch_billing())
            except Exception as e:
                QtWidgets.QMessageBox.critical(preview_window, "Error", f"Failed to issue billing: {e}")

    def open_quick_issue_dialog(self):
        dialog = QuickIssueDialog(parent=self)
        dialog.exec_()

    def open_quick_set_paid_dialog(self):
        dialog = QuickSetPaidDialog(parent=self)
        dialog.exec_()

    def confirm_void_bill(self, preview_window, billing_data):
        reply = QtWidgets.QMessageBox.question(
            preview_window,
            "Confirm Void",
            f"Are you sure you want to void billing {billing_data[0]} and its associated reading and transaction?",
            QtWidgets.QMessageBox.Yes | QtWidgets.QMessageBox.No
        )
        if reply == QtWidgets.QMessageBox.Yes:
            try:
                backend = adminPageBack()
                billing_code = billing_data[0]
                billing_id = backend.get_billing_id(billing_code)

                # Mark billing as VOID
                backend.update_billing_status(billing_id, "VOID")

                # Get reading_id and transaction_id
                reading_id = backend.get_reading_id_by_billing_id(billing_id)
                trans_id = backend.get_transaction_id_by_billing_id(billing_id)

                # Mark transaction as voided
                if trans_id:
                    backend.update_transaction_status(trans_id, "VOID")

                # Mark reading as voided and update meter with previous reading
                if reading_id:
                    reading_info = backend.get_reading_info_by_id(reading_id)
                    if reading_info:
                        reading_prev = reading_info[2]      # reading_prev
                        reading_date = reading_info[1]     # reading_date
                        meter_id = reading_info[4]         # meter_id

                        backend.void_reading(reading_id)

                        # Find matching non-voided reading with current = voided reading_prev
                        matching_reading = backend.get_reading_by_current_and_meter(reading_prev, meter_id)

                        print("matching_reading result:", matching_reading)
                        print("type:", type(matching_reading))

                        if matching_reading:
                            # reading_current matched, use its date
                            updated_reading_date = matching_reading[1]
                        else:
                            # fallback: use current reading_date from voided reading
                            updated_reading_date = reading_date

                        # Update the meter's last reading and date
                        backend.update_meter_latest_reading(reading_prev, updated_reading_date, meter_id)

                QtWidgets.QMessageBox.information(preview_window, "Voided", f"Billing {billing_code} and its related data voided successfully.")
                preview_window.close()
                self.populate_table(backend.fetch_billing())
            except Exception as e:
                QtWidgets.QMessageBox.critical(preview_window, "Error", f"Failed to void billing and related data: {e}")




    

    def on_solo_generation_finished(self, merged_pdf_path):
        self.loading.close()
        self.preview_window = self.BillPreview(merged_pdf_path)

        # 🔧 FIX: Get the correct billing data used
        billing_data = self.last_printed_bill_data  # Store this when you call print_bill()

        # Mark as Printed button
        mark_btn = QtWidgets.QPushButton("✅ Mark as Printed")
        mark_btn.setStyleSheet("""
            QPushButton {
                background-color: #43A047;
                color: white;
                padding: 6px 12px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2E7D32;
            }
        """)
        mark_btn.clicked.connect(lambda: self.mark_as_printed(self.preview_window, [billing_data]))
        
        self.preview_window.layout().addWidget(mark_btn)
        self.preview_window.show()

    def preview_generated_pdf(self, pdf_path, batch_mode=True):
        # Display the PDF preview window using BillPreview class
        self.preview_window = self.BillPreview(pdf_path, batch_mode=batch_mode)
        self.preview_window.show()

    class BillPreview(QWidget):
        def __init__(self, pdf_path, batch_mode=False):
            super().__init__()
            self.pdf_path = pdf_path
            self.batch_mode = batch_mode

            self.setWindowTitle("Bill Preview")
            self.setMinimumSize(900, 700)

            # Scrollable layout
            scroll = QScrollArea(self)
            layout = QVBoxLayout(self)
            layout.addWidget(scroll)

            container = QWidget()
            scroll.setWidget(container)
            scroll.setWidgetResizable(True)

            container_layout = QVBoxLayout(container)

            try:
                # Load PDF pages as images using PyMuPDF (fitz)
                try:
                    doc = fitz.open(pdf_path)
                except Exception as e:
                    error_label = QLabel(f"❌ Failed to open PDF: {e}")
                    container_layout.addWidget(error_label)
                    return
                for i, page in enumerate(doc):
                    try:
                        # Render page to image
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        image_format = QImage.Format_RGBA8888 if pix.alpha else QImage.Format_RGB888
                        qimage = QImage(pix.samples, pix.width, pix.height, pix.stride, image_format)
                        pixmap = QPixmap.fromImage(qimage)

                        # Add label for batch mode
                        if self.batch_mode:
                            page_label = QLabel(f"📄 - BILL {i + 1}")
                            page_label.setAlignment(QtCore.Qt.AlignCenter)
                            page_label.setStyleSheet("color: red; font-size: 14px; font-weight: bold; padding: 10px;")
                            container_layout.addWidget(page_label)

                        # Show rendered image
                        image_label = QLabel()
                        image_label.setPixmap(pixmap)
                        image_label.setAlignment(QtCore.Qt.AlignCenter)
                        container_layout.addWidget(image_label)

                        # Reprint button (batch mode only)
                        if self.batch_mode:
                            reprint_button = QPushButton(f"🖨 Reprint BILL {i + 1}")
                            reprint_button.setStyleSheet("""
                                QPushButton {
                                    background-color: #4CAF50;
                                    color: white;
                                    font-weight: bold;
                                    padding: 6px 12px;
                                    border-radius: 4px;
                                }
                                QPushButton:hover {
                                    background-color: #45A049;
                                }
                            """)
                            reprint_button.clicked.connect(lambda _, idx=i: self.print_single_page(idx))
                            container_layout.addWidget(reprint_button)

                    except Exception as page_error:
                        error_label = QLabel(f"❌ Failed to load page {i + 1}: {page_error}")
                        container_layout.addWidget(error_label)

                    # Spacing between each bill
                    spacer = QLabel()
                    spacer.setFixedHeight(40)
                    container_layout.addWidget(spacer)

            except Exception as e:
                error_label = QLabel(f"❌ Failed to load preview:\n{str(e)}")
                error_label.setAlignment(QtCore.Qt.AlignCenter)
                error_label.setStyleSheet("color: red; font-weight: bold; padding: 20px;")
                container_layout.addWidget(error_label)

            # Print button
            print_btn = QPushButton("Print")
            print_btn.clicked.connect(self.print_pdf)
            layout.addWidget(print_btn)

        def print_pdf(self):
            printer = QPrinter()
            printer.setPageSize(QPrinter.A4)
            printer.setFullPage(True)

            # Show print dialog
            dialog = QPrintDialog(printer, self)
            if dialog.exec_() == QPrintDialog.Accepted:
                try:
                    import fitz  # PyMuPDF
                    doc = fitz.open(self.pdf_path)
                    painter = QtGui.QPainter(printer)

                    for i, page in enumerate(doc):
                        if i > 0:
                            printer.newPage()

                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        image_format = QImage.Format_RGBA8888 if pix.alpha else QImage.Format_RGB888
                        qimage = QImage(pix.samples, pix.width, pix.height, pix.stride, image_format)
                        pixmap = QPixmap.fromImage(qimage)

                        rect = painter.viewport()
                        size = pixmap.size()
                        size.scale(rect.size(), QtCore.Qt.KeepAspectRatio)
                        painter.setViewport(rect.x(), rect.y(), size.width(), size.height())
                        painter.setWindow(pixmap.rect())
                        painter.drawPixmap(0, 0, pixmap)

                    painter.end()
                except Exception as e:
                    QtWidgets.QMessageBox.warning(self, "Error", f"Failed to print PDF: {str(e)}")
        
        def print_single_page(self, page_index):
            printer = QPrinter()
            printer.setPageSize(QPrinter.A4)
            printer.setFullPage(True)

            dialog = QPrintDialog(printer, self)
            if dialog.exec_() == QPrintDialog.Accepted:
                try:
                    import fitz
                    doc = fitz.open(self.pdf_path)
                    page = doc[page_index]

                    painter = QtGui.QPainter(printer)

                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    image_format = QImage.Format_RGBA8888 if pix.alpha else QImage.Format_RGB888
                    qimage = QImage(pix.samples, pix.width, pix.height, pix.stride, image_format)
                    pixmap = QPixmap.fromImage(qimage)

                    rect = painter.viewport()
                    size = pixmap.size()
                    size.scale(rect.size(), QtCore.Qt.KeepAspectRatio)
                    painter.setViewport(rect.x(), rect.y(), size.width(), size.height())
                    painter.setWindow(pixmap.rect())
                    painter.drawPixmap(0, 0, pixmap)

                    painter.end()
                except Exception as e:
                    QtWidgets.QMessageBox.warning(self, "Error", f"Failed to reprint: {str(e)}")


    def print_all_on_page(self):
        from PyQt5.QtPrintSupport import QPrinter, QPrintDialog
        import os

        # Base path setup
        base_dir = os.path.abspath(os.path.join(os.path.dirname(sys.argv[0]), ".."))
        template_path = os.path.join(base_dir, "templates", "bill_template.docx")

        if not os.path.exists(template_path):
            QtWidgets.QMessageBox.warning(self, "Template Missing", f"Template not found at:\n{template_path}")
            return

        # Prepare temp folder
        temp_folder = os.path.join(base_dir, "temp_bills")
        os.makedirs(temp_folder, exist_ok=True)

        # Clear old files
        for f in os.listdir(temp_folder):
            file_path = os.path.join(temp_folder, f)
            if os.path.isfile(file_path):
                os.unlink(file_path)


        # Get current page data
        start_idx = (self.current_page - 1) * self.rows_per_page
        end_idx = min(start_idx + self.rows_per_page, len(self.filtered_data))
        page_data = self.filtered_data[start_idx:end_idx]

        # Show animated loading dialog
        self.loading = self.LoadingDialog("Generating bills...")
        self.loading.set_progress(0, len(page_data))
        self.loading.show()

        # Start background generation
        worker = BillWorker(page_data, template_path, temp_folder, self.fill_word_template, self.convert_to_pdf)
        worker.signals.progress.connect(lambda i, total: (
            self.loading.set_message(f"Generating BILL {i} of {total}"),
            self.loading.set_progress(i, total)
        ))
        worker.signals.finished.connect(self.on_generation_finished)
        worker.signals.error.connect(self.on_generation_failed)

        QThreadPool.globalInstance().start(worker)



    def on_generation_finished(self, merged_pdf_path):
        self.loading.close()
        self.preview_window = self.BillPreview(merged_pdf_path, batch_mode=True)

        # Mark All as Printed button
        mark_all_btn = QtWidgets.QPushButton("✅ Mark All as Printed")
        mark_all_btn.setStyleSheet("""
            QPushButton {
                background-color: #388E3C;
                color: white;
                padding: 6px 12px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2E7D32;
            }
        """)
        start_idx = (self.current_page - 1) * self.rows_per_page
        end_idx = min(start_idx + self.rows_per_page, len(self.filtered_data))
        page_data = self.filtered_data[start_idx:end_idx]
        mark_all_btn.clicked.connect(lambda: self.mark_as_printed(self.preview_window, page_data))

        self.preview_window.layout().addWidget(mark_all_btn)
        self.preview_window.show()

    def on_generation_failed(self, error_msg):
        self.loading.close()
        QtWidgets.QMessageBox.critical(self, "Generation Failed", error_msg)




    def paint_bill(self, printer, billing_data):
        """Paint the bill content for printing"""
        painter = QtGui.QPainter(printer)
        
        if not painter.isActive():
            print("Painter is not active")
            return
        
        try:
            # Get billing data
            billing_code, issued_date, billing_due, client_name, client_lname, client_location, billing_total, status = billing_data
            
            # Set up fonts
            title_font = QtGui.QFont("Arial", 18, QtGui.QFont.Bold)
            header_font = QtGui.QFont("Arial", 14, QtGui.QFont.Bold)
            subheader_font = QtGui.QFont("Arial", 12, QtGui.QFont.Bold)
            content_font = QtGui.QFont("Arial", 11)
            small_font = QtGui.QFont("Arial", 9)
            
            # Get page dimensions
            page_rect = printer.pageRect()
            margin = 80
            content_width = page_rect.width() - (2 * margin)
            
            # Starting position
            y_pos = margin
            line_height = 30
            section_spacing = 40
            
            # Helper function to draw centered text
            def draw_centered_text(text, font, y_position, height=30):
                painter.setFont(font)
                text_rect = QtCore.QRect(margin, y_position, content_width, height)
                painter.drawText(text_rect, QtCore.Qt.AlignCenter, text)
                return y_position + height
            
            # Helper function to draw left-aligned text
            def draw_left_text(text, font, y_position, height=25):
                painter.setFont(font)
                text_rect = QtCore.QRect(margin, y_position, content_width, height)
                painter.drawText(text_rect, QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter, text)
                return y_position + height
            
            # Helper function to draw right-aligned text
            def draw_right_text(text, font, y_position, height=25):
                painter.setFont(font)
                text_rect = QtCore.QRect(margin, y_position, content_width, height)
                painter.drawText(text_rect, QtCore.Qt.AlignRight | QtCore.Qt.AlignVCenter, text)
                return y_position + height
            
            # Title
            y_pos = draw_centered_text("WATER BILLING STATEMENT", title_font, y_pos, 40)
            y_pos += 20
            
            # Company info
            y_pos = draw_centered_text("WATER UTILITY COMPANY", header_font, y_pos, 35)
            y_pos += section_spacing
            
            # Draw a horizontal line
            painter.drawLine(margin, y_pos, margin + content_width, y_pos)
            y_pos += 20
            
            # Bill Information Section
            y_pos = draw_left_text("BILL INFORMATION", subheader_font, y_pos, 30)
            y_pos += 10
            
            painter.setFont(content_font)
            
            # Create two columns for bill info
            left_column_x = margin
            right_column_x = margin + (content_width // 2)
            
            # Left column
            current_y = y_pos
            painter.drawText(left_column_x, current_y, f"Billing Code: {billing_code}")
            current_y += line_height
            painter.drawText(left_column_x, current_y, f"Issue Date: {issued_date}")
            current_y += line_height
            
            # Right column
            current_y = y_pos
            painter.drawText(right_column_x, current_y, f"Due Date: {billing_due}")
            current_y += line_height
            painter.drawText(right_column_x, current_y, f"Status: {status}")
            current_y += line_height
            
            y_pos = current_y + 20
            
            # Customer Information Section
            y_pos = draw_left_text("CUSTOMER INFORMATION", subheader_font, y_pos, 30)
            y_pos += 10
            
            painter.setFont(content_font)
            painter.drawText(margin, y_pos, f"Name: {client_name} {client_lname}")
            y_pos += line_height
            painter.drawText(margin, y_pos, f"Location: {client_location}")
            y_pos += section_spacing
            
            # Billing Details Section
            y_pos = draw_left_text("BILLING DETAILS", subheader_font, y_pos, 30)
            y_pos += 10
            
            # Create table
            table_top = y_pos
            table_height = 80
            table_rect = QtCore.QRect(margin, table_top, content_width, table_height)
            
            # Draw table border
            painter.drawRect(table_rect)
            
            # Draw table header
            header_height = 25
            painter.setFont(subheader_font)
            
            # Table headers
            desc_rect = QtCore.QRect(margin + 10, table_top + 5, content_width - 200, header_height)
            amount_rect = QtCore.QRect(margin + content_width - 180, table_top + 5, 170, header_height)
            
            painter.drawText(desc_rect, QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter, "Description")
            painter.drawText(amount_rect, QtCore.Qt.AlignRight | QtCore.Qt.AlignVCenter, "Amount")
            
            # Draw horizontal line under header
            header_line_y = table_top + header_height + 5
            painter.drawLine(margin, header_line_y, margin + content_width, header_line_y)
            
            # Table content
            painter.setFont(content_font)
            content_y = header_line_y + 10
            
            desc_content_rect = QtCore.QRect(margin + 10, content_y, content_width - 200, 20)
            amount_content_rect = QtCore.QRect(margin + content_width - 180, content_y, 170, 20)
            
            painter.drawText(desc_content_rect, QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter, "Water Consumption")
            painter.drawText(amount_content_rect, QtCore.Qt.AlignRight | QtCore.Qt.AlignVCenter, f"₱{billing_total}")
            
            y_pos = table_top + table_height + 30
            
            # Total Section
            painter.setFont(header_font)
            total_rect = QtCore.QRect(margin + content_width - 250, y_pos, 240, 30)
            painter.drawText(total_rect, QtCore.Qt.AlignRight | QtCore.Qt.AlignVCenter, f"TOTAL AMOUNT: ₱{billing_total}")
            
            # Footer
            footer_y = page_rect.height() - margin - 60
            painter.setFont(small_font)
            
            # Draw footer line
            painter.drawLine(margin, footer_y - 10, margin + content_width, footer_y - 10)
            
            footer_rect = QtCore.QRect(margin, footer_y, content_width, 30)
            painter.drawText(footer_rect, QtCore.Qt.AlignCenter, "Thank you for your payment!")
            
            # Add payment instructions
            payment_y = footer_y + 25
            payment_rect = QtCore.QRect(margin, payment_y, content_width, 20)
            painter.drawText(payment_rect, QtCore.Qt.AlignCenter, "Please pay on or before the due date to avoid penalties.")
            
        except Exception as e:
            print(f"Error painting bill: {str(e)}")
            # Draw error message on the page
            painter.setFont(QtGui.QFont("Arial", 12))
            error_rect = QtCore.QRect(margin, margin, content_width, 50)
            painter.drawText(error_rect, QtCore.Qt.AlignCenter, f"Error generating bill: {str(e)}")
        
        # Note: Don't call painter.end() here - QPrintPreviewDialog handles it


    def mark_as_printed(self, preview_window, bills_to_update):
        backend = adminPageBack()
        success = 0
        for bill in bills_to_update:
            billing_code = bill[0]
            billing_id = backend.get_billing_id(billing_code)
            try:
                backend.update_billing_status(billing_id, "PRINTED")
                success += 1
            except Exception as e:
                print(f"❌ Failed to update bill {billing_code}: {e}")
        QtWidgets.QMessageBox.information(self, "Update Complete", f"{success} bill(s) marked as PRINTED.")
        preview_window.close()
        self.populate_table(backend.fetch_billing())


    def setup_ui(self):
        layout = QtWidgets.QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)

        # Define input style at the beginning
        input_style = """
            QLineEdit, QComboBox {
                padding: 8px;
                border: 1px solid #ccc;
                border-radius: 4px;
                font-family: 'Roboto', sans-serif;
                min-width: 250px;
            }
        """

        # Header with title and search
        header_layout = QtWidgets.QHBoxLayout()
        title = QtWidgets.QLabel("BILLING LIST")
        title.setStyleSheet("""
            font-family: 'Montserrat', sans-serif;
            font-size: 24px;
            font-weight: bold;
        """)
        header_layout.addWidget(title)
        header_layout.addStretch()

        # Search and Add button container
        search_add_layout = QtWidgets.QHBoxLayout() 
        
        # Search container
        search_container = QtWidgets.QHBoxLayout()
        
        # Search criteria dropdown
        self.search_criteria = QtWidgets.QComboBox()
        self.search_criteria.addItems(["BILLING CODE", "CLIENT NAME", "CLIENT LOCATION"])
        self.search_criteria.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 1px solid #ccc;
                border-radius: 4px;
                font-family: 'Roboto', sans-serif;
                min-width: 120px;
                background-color: white;
            }
        """)
        search_container.addWidget(self.search_criteria)
        
        # Search input
        self.search_input = QtWidgets.QLineEdit()
        self.search_input.setPlaceholderText("Search billings...")
        self.search_input.setStyleSheet(input_style)
        search_container.addWidget(self.search_input)
        self.search_input.textChanged.connect(self.filter_table)
        
        # Add status filter dropdown
        self.status_filter = QtWidgets.QComboBox()
        self.status_filter.addItems(["ALL", "TO BE PRINTED", "PRINTED", "VOID", "PENDING PAYMENT", "PAID"])
        self.status_filter.setStyleSheet("""
            QComboBox {
                padding: 8px;
                border: 1px solid #ccc;
                border-radius: 4px;
                font-family: 'Roboto', sans-serif;
                min-width: 150px;
                background-color: white;
            }
        """)
        self.status_filter.currentTextChanged.connect(self.filter_table)
        search_container.addWidget(self.status_filter)
        
        # Add search container to search_add_layout
        search_add_layout.addLayout(search_container)
        
        # Add button with icon
        add_btn = QtWidgets.QPushButton("ADD BILLING", icon=QtGui.QIcon("../images/add.png"))
        add_btn.setStyleSheet("""
            QPushButton {
                background-color: rgb(229, 115, 115);
                color: white;
                padding: 8px 15px;
                border-radius: 4px;
                font-family: 'Roboto', sans-serif;
            }
            QPushButton:hover {
                background-color: rgb(200, 100, 100);
            }
        """)
        add_btn.clicked.connect(self.show_add_billing)
        search_add_layout.addWidget(add_btn)

        # Quick Set Paid Button
        quick_paid_btn = QtWidgets.QPushButton("✅ Quick Set Paid")
        quick_paid_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                padding: 8px 15px;
                border-radius: 4px;
                font-family: 'Roboto', sans-serif;
                min-width: 140px;
            }
            QPushButton:hover {
                background-color: #388E3C;
            }
        """)
        quick_paid_btn.clicked.connect(self.open_quick_set_paid_dialog)

        search_add_layout.addWidget(quick_paid_btn)

        # Quick Issue Billing Button
        quick_issue_btn = QtWidgets.QPushButton("⚡ Quick Issue Billing")
        quick_issue_btn.setStyleSheet("""
            QPushButton {
                background-color: #FFA726;
                color: white;
                padding: 8px 15px;
                border-radius: 4px;
                font-family: 'Roboto', sans-serif;
                min-width: 140px;
            }
            QPushButton:hover {
                background-color: #FB8C00;
            }
        """)
        quick_issue_btn.clicked.connect(self.open_quick_issue_dialog)

        search_add_layout.addWidget(quick_issue_btn)
        
        # Add search_add_layout to header_layout
        header_layout.addLayout(search_add_layout)
        layout.addLayout(header_layout)
        
        # Create billing table before accessing it
        self.billing_table = QtWidgets.QTableWidget()
        self.billing_table.setAlternatingRowColors(True)
        self.billing_table.setStyleSheet("""
            QTableWidget {
                border: 1px solid #ccc;
                border-radius: 4px;
                background-color: #E8F5E9;
                alternate-background-color: #FFFFFF;
            }
            QHeaderView::section {
                background-color: #B2C8B2;
                padding: 8px;
                border: none;
                font-family: 'Roboto', sans-serif;
                font-weight: bold;
            }
        """)
        
        # Set up table columns - Added ACTION column
        self.billing_table.setColumnCount(9)  # Increased from 8 to 9
        self.billing_table.setHorizontalHeaderLabels([
            "BILLING CODE", "ISSUED DATE", "BILLING DUE", 
            "NAME", "LAST NAME","LOCATION", "BILLING TOTAL", "STATUS", "ACTION"
        ])

        # Set the table to fill all available space
        self.billing_table.setSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Expanding)
        self.billing_table.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.Stretch)
        
        # Set ACTION column to fixed width
        self.billing_table.horizontalHeader().setSectionResizeMode(8, QtWidgets.QHeaderView.Fixed)
        self.billing_table.setColumnWidth(8, 350)  # Adjust based on your needs
        
        # Enable horizontal scrollbar
        self.billing_table.setHorizontalScrollMode(QtWidgets.QAbstractItemView.ScrollPerPixel)
        self.billing_table.setWordWrap(False)

        billing_back = adminPageBack()
        self.all_data = billing_back.fetch_billing()
        self.filtered_data = self.all_data.copy()  # Initialize filtered data with all data
        
        # Now we can safely adjust table properties
        self.billing_table.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.Stretch)
        # Reset ACTION column to fixed width after stretch
        self.billing_table.horizontalHeader().setSectionResizeMode(8, QtWidgets.QHeaderView.Fixed)
        self.billing_table.setColumnWidth(8, 350)
        
        self.billing_table.setSelectionBehavior(QtWidgets.QTableWidget.SelectRows)
        self.billing_table.verticalHeader().setVisible(False)
        self.billing_table.setEditTriggers(QtWidgets.QTableWidget.NoEditTriggers)
        
        # Create a custom delegate for text elision with tooltip
        delegate = TextEllipsisDelegate(self.billing_table)
        self.billing_table.setItemDelegate(delegate)

        # Add table to the main layout with full expansion
        layout.addWidget(self.billing_table)
        
        # Add pagination controls
        pagination_layout = QtWidgets.QHBoxLayout()
        pagination_layout.setAlignment(QtCore.Qt.AlignCenter)
        
        # First page button
        self.first_page_btn = QtWidgets.QPushButton("⏮ First")
        self.first_page_btn.clicked.connect(self.go_to_first_page)
        
        # Previous page button
        self.prev_page_btn = QtWidgets.QPushButton("◀ Previous")
        self.prev_page_btn.clicked.connect(self.go_to_prev_page)
        
        # Page indicator
        self.page_indicator = QtWidgets.QLabel("Page 1 of 1")
        self.page_indicator.setAlignment(QtCore.Qt.AlignCenter)
        self.page_indicator.setStyleSheet("font-weight: bold; min-width: 150px;")
        
        # Next page button
        self.next_page_btn = QtWidgets.QPushButton("Next ▶")
        self.next_page_btn.clicked.connect(self.go_to_next_page)
        
        # Last page button
        self.last_page_btn = QtWidgets.QPushButton("Last ⏭")
        self.last_page_btn.clicked.connect(self.go_to_last_page)
        
        # Records per page selector
        self.page_size_label = QtWidgets.QLabel("Records per page:")
        self.rows_per_page_combo = QtWidgets.QComboBox()
        self.rows_per_page_combo.addItems(["5", "10", "20", "50", "100"])
        self.rows_per_page_combo.setCurrentText(str(self.rows_per_page))
        self.rows_per_page_combo.currentTextChanged.connect(self.change_rows_per_page)
        
        # Style for pagination buttons
        pagination_btn_style = """
            QPushButton {
                background-color: #81C784;
                color: white;
                padding: 5px 10px;
                border-radius: 4px;
                min-width: 80px;
            }
            QPushButton:hover {
                background-color: #66BB6A;
            }
            QPushButton:disabled {
                background-color: #E0E0E0;
                color: #9E9E9E;
            }
        """
        
        self.first_page_btn.setStyleSheet(pagination_btn_style)
        self.prev_page_btn.setStyleSheet(pagination_btn_style)
        self.next_page_btn.setStyleSheet(pagination_btn_style)
        self.last_page_btn.setStyleSheet(pagination_btn_style)
        
        # Add pagination controls to layout
        pagination_layout.addWidget(self.first_page_btn)
        pagination_layout.addWidget(self.prev_page_btn)
        pagination_layout.addWidget(self.page_indicator)
        pagination_layout.addWidget(self.next_page_btn)
        pagination_layout.addWidget(self.last_page_btn)
        pagination_layout.addSpacing(20)
        pagination_layout.addWidget(self.page_size_label)
        pagination_layout.addWidget(self.rows_per_page_combo)
        
        layout.addLayout(pagination_layout)

        # Add batch print button
        self.batch_print_btn = QtWidgets.QPushButton("🖨 Print All on This Page")
        self.update_batch_print_button_state()
        self.batch_print_btn.setStyleSheet("""
            QPushButton {
                background-color: #2196F3;
                color: white;
                padding: 8px 16px;
                border-radius: 4px;
            }
            QPushButton:hover:enabled {
                background-color: #1976D2;
            }
            QPushButton:disabled {
                background-color: #cccccc;
                color: #666666;
                border: 1px solid #aaaaaa;
            }
        """)

        self.batch_print_btn.clicked.connect(self.print_all_on_page)
        layout.addWidget(self.batch_print_btn)

        
        # Calculate total pages and update table
        self.update_pagination()

    def update_batch_print_button_state(self):
        # Disable the batch print button when the status is "ALL" or not "TO BE PRINTED"
        if self.selected_status != "TO BE PRINTED":
            self.batch_print_btn.setEnabled(False)
            self.batch_print_btn.setToolTip("Batch printing is only available for bills marked 'TO BE PRINTED'.")
        else:
            self.batch_print_btn.setEnabled(True)
            self.batch_print_btn.setToolTip("")  # Reset the tooltip



    def update_pagination(self):
        # Calculate total pages based on filtered data
        visible_rows = len(self.filtered_data)
        self.total_pages = max(1, math.ceil(visible_rows / self.rows_per_page))
        
        # Adjust current page if it's beyond the new total
        if self.current_page > self.total_pages:
            self.current_page = self.total_pages
        
        # Update page indicator
        self.page_indicator.setText(f"Page {self.current_page} of {self.total_pages}")
        
        # Enable/disable navigation buttons
        self.first_page_btn.setEnabled(self.current_page > 1)
        self.prev_page_btn.setEnabled(self.current_page > 1)
        self.next_page_btn.setEnabled(self.current_page < self.total_pages)
        self.last_page_btn.setEnabled(self.current_page < self.total_pages)
        
        # Update table with current page data
        self.populate_table_for_page()
        self.update_batch_print_button_state()

    def populate_table_for_page(self):
        data_to_display = self.filtered_data
        
        # Calculate start and end indices for the current page
        start_idx = (self.current_page - 1) * self.rows_per_page
        end_idx = min(start_idx + self.rows_per_page, len(data_to_display))
        
        # Get data for current page
        page_data = data_to_display[start_idx:end_idx]
        
        # Set the table row count
        self.billing_table.setRowCount(len(page_data))
        
        # Fill the table with data
        for i, billing in enumerate(page_data):
            # Unpack values
            billing_code, issued_date, billing_due, client_name, client_lname,  client_location, billing_total, status = billing
            
            # Add billing data to the table
            self.create_scrollable_cell(i, 0, str(billing_code))
            self.create_scrollable_cell(i, 1, str(issued_date))
            self.create_scrollable_cell(i, 2, str(billing_due))
            self.create_scrollable_cell(i, 3, client_name)
            self.create_scrollable_cell(i, 4, client_lname)
            self.create_scrollable_cell(i, 5, client_location)
            self.create_scrollable_cell(i, 6, str(billing_total))
            
            # Status with color coding
            status_item = QtWidgets.QTableWidgetItem(status)
            status_item.setForeground(
                QtGui.QColor("#64B5F6") if status == "PAID" else QtGui.QColor("#E57373")
            )
            self.billing_table.setItem(i, 7, status_item)
            
            # Add action cell with print button
            self.create_action_cell(i, billing)

    def populate_table(self, data):
        # Update all data
        self.all_data = data
        self.filtered_data = data.copy()

        # Reapply current filter and search
        self.filter_table()
        

    def filter_table(self):
        search_by = self.search_criteria.currentText()
        search_text = self.search_input.text().lower()
        status_filter = self.status_filter.currentText()

        column_mapping = {
            "CLIENT NAME": 3,
            "CLIENT LAST NAME": 4,
            "CLIENT LOCATION": 5
        }

        col_index = column_mapping.get(search_by, 0)

        self.filtered_data = self.all_data.copy()

        # Apply search filter
        if search_text:
            self.filtered_data = [
                row for row in self.filtered_data 
                if search_text in str(row[col_index]).lower()
            ]

        # Apply status filter
        if status_filter != "ALL":
            self.filtered_data = [
                row for row in self.filtered_data
                if row[7] == status_filter
            ]

        # Update selected status for button logic
        self.selected_status = status_filter  

        # Update pagination
        self.current_page = 1
        self.update_pagination()

        self.update_batch_print_button_state()


    def change_rows_per_page(self, value):
        self.rows_per_page = int(value)
        self.current_page = 1  # Reset to first page
        self.update_pagination()
    
    def go_to_first_page(self):
        if self.current_page != 1:
            self.current_page = 1
            self.update_pagination()
    
    def go_to_prev_page(self):
        if self.current_page > 1:
            self.current_page -= 1
            self.update_pagination()
    
    def go_to_next_page(self):
        if self.current_page < self.total_pages:
            self.current_page += 1
            self.update_pagination()
    
    def go_to_last_page(self):
        if self.current_page != self.total_pages:
            self.current_page = self.total_pages
            self.update_pagination()

    def show_add_billing(self):
        dialog = QtWidgets.QDialog(self)
        dialog.setWindowTitle("Add New Bill")
        dialog.setFixedSize(1000, 700)
        dialog.setStyleSheet("""
            QDialog {
                background-color: #C9EBCB;
            }
            QLabel {
                font-family: 'Arial', sans-serif;
                font-weight: bold;
            }
        """)
        
        layout = QtWidgets.QVBoxLayout(dialog)
        layout.setContentsMargins(30, 5, 30, 5)
        layout.setSpacing(10)

        # Section Title
        title = QtWidgets.QLabel("BILLING INFORMATION FORM")
        title.setStyleSheet("""
            font-size: 20px;
            padding: 10px;
        """)
        title.setAlignment(QtCore.Qt.AlignCenter)
        layout.addWidget(title)

        # Form Layout
        form_layout = QtWidgets.QGridLayout()
        form_layout.setHorizontalSpacing(40)
        form_layout.setVerticalSpacing(20)

        input_style = """
            QLineEdit, QDateEdit, QComboBox {
                font-family: 'Arial';
                font-size: 14px;
                padding: 8px;
                border: 1px solid #bdc3c7;
                border-radius: 4px;
                background-color: #ffffff;
            }
        """

        readonly_style = """
        QLineEdit {
            font-family: 'Arial';
            font-size: 14px;
            padding: 8px;
            border: 1px solid #bdc3c7;
            border-radius: 4px;
            background-color: #e0e0e0;
            color: #555555;
        }
    """


        # --- LEFT COLUMN ---

        def create_labeled_widget(label_text, widget):
            layout = QtWidgets.QVBoxLayout()
            label = QtWidgets.QLabel(label_text)
            label.setFont(QtGui.QFont("Arial", 10))
            layout.addWidget(label)
            layout.addWidget(widget)
            return layout

        
        client = QtWidgets.QComboBox()
        client.setStyleSheet(input_style)

        reading_date = QtWidgets.QDateEdit()
        reading_date.setCalendarPopup(True)
        reading_date.setStyleSheet(input_style)
        reading_date.setMaximumDate(QtCore.QDate.currentDate())
        reading_date.setDate(QtCore.QDate.currentDate())  # Set current date as default
         

        previous_reading = QtWidgets.QLineEdit()
        previous_reading.setReadOnly(True)
        previous_reading.setStyleSheet(readonly_style)

        present_reading = QtWidgets.QLineEdit()
        present_reading.setEnabled(False)
        present_reading.setStyleSheet(readonly_style)

        cubic_meter_consumed = QtWidgets.QLineEdit()
        cubic_meter_consumed.setReadOnly(True)
        cubic_meter_consumed.setStyleSheet(readonly_style)

        amount = QtWidgets.QLineEdit()
        amount.setReadOnly(True)
        amount.setStyleSheet(readonly_style)

        due_date = QtWidgets.QDateEdit()
        due_date.setCalendarPopup(True)
        due_date.setStyleSheet(input_style)
        due_date.setMinimumDate(QtCore.QDate.currentDate())

        # Bold centered section header
        additional_charge_label = QtWidgets.QLabel("ADDITIONAL CHARGE")
        font = QtGui.QFont()
        font.setBold(True)
        font.setPointSize(11)
        additional_charge_label.setFont(font)
        additional_charge_label.setAlignment(QtCore.Qt.AlignCenter)
        form_layout.addWidget(additional_charge_label, 0, 1)

        subscribe_capital = QtWidgets.QLineEdit()
        subscribe_capital.setStyleSheet(input_style)

        late_payment = QtWidgets.QLineEdit()
        late_payment.setStyleSheet(input_style)

        penalty = QtWidgets.QLineEdit()
        penalty.setStyleSheet(input_style)

        total_charge = QtWidgets.QLineEdit()
        total_charge.setStyleSheet(input_style)
        total_charge.setReadOnly(True)
        total_charge.setStyleSheet(readonly_style)

        total_bill = QtWidgets.QLineEdit()
        total_bill.setStyleSheet(input_style)
        total_bill.setReadOnly(True)
        total_bill.setStyleSheet(readonly_style)

        form_layout.addLayout(create_labeled_widget("SUBSCRIBE CAPITAL:", subscribe_capital), 1, 1)
        form_layout.addLayout(create_labeled_widget("LATE PAYMENT:", late_payment), 2, 1)
        form_layout.addLayout(create_labeled_widget("PENALTY:", penalty), 3, 1)
        form_layout.addLayout(create_labeled_widget("TOTAL CHARGE:", total_charge), 4, 1)
        form_layout.addLayout(create_labeled_widget("TOTAL BILL:", total_bill), 6, 1)

        # Add form_layout to main layout
        layout.addLayout(form_layout)


        # Button Container
        button_container = QtWidgets.QWidget()
        button_layout = QtWidgets.QHBoxLayout(button_container)
        button_layout.setAlignment(QtCore.Qt.AlignRight)

        # Cancel Button
        cancel_btn = QtWidgets.QPushButton("Cancel")
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                padding: 8px 15px;
                border-radius: 4px;
                font-family: 'Roboto', sans-serif;
                min-width: 100px;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)

        # Save Button
        save_btn = QtWidgets.QPushButton("Save Bill")
        save_btn.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                padding: 8px 15px;
                border-radius: 4px;
                font-family: 'Roboto', sans-serif;
                min-width: 100px;
            }
            QPushButton:hover {
                background-color: #219a52;
            }
        """)

        button_layout.addWidget(cancel_btn)
        button_layout.addWidget(save_btn)
        layout.addWidget(button_container)
        


        form_layout.addLayout(create_labeled_widget("CLIENT:", client), 0, 0)
        form_layout.addLayout(create_labeled_widget("READING DATE:", reading_date), 1, 0)
        form_layout.addLayout(create_labeled_widget("PREVIOUS READING:", previous_reading), 2, 0)
        form_layout.addLayout(create_labeled_widget("PRESENT READING:", present_reading), 3, 0)
        form_layout.addLayout(create_labeled_widget("CUBIC METER CONSUMED:", cubic_meter_consumed), 4, 0)
        form_layout.addLayout(create_labeled_widget("AMOUNT:", amount), 5, 0)
        form_layout.addLayout(create_labeled_widget("DUE DATE:", due_date), 6, 0)

        IadminPageBack = adminPageBack()
        client.setEditable(True)
        client.setInsertPolicy(QtWidgets.QComboBox.NoInsert)
        client.setStyleSheet(input_style)
        client.lineEdit().setReadOnly(False)
        client.clear()

        # Populate client ComboBox with active clients only
        clients = IadminPageBack.fetch_active_clients()  # Use new method
        client_entries = []
        client_data_map = {}

        for client_data in clients:
            client_id = client_data[0]
            client_number = client_data[1]
            first_name = client_data[2]
            last_name = client_data[4]
            display_text = f"{client_number} - {first_name} {last_name}"
            client.addItem(display_text, client_id)
            client_entries.append(display_text)
            client_data_map[display_text] = client_id

        #Completer for client ComboBox
        completer = QtWidgets.QCompleter(client_entries)
        completer.setCaseSensitivity(QtCore.Qt.CaseInsensitive)
        completer.setFilterMode(QtCore.Qt.MatchContains)
        client.setCompleter(completer)

        client.setCurrentText("Select Client")

        def update_total_bill():
            try:
                amt_text = amount.text().strip()
                charge_text = total_charge.text().strip()

                amt = float(amt_text) if amt_text else 0
                charge = float(charge_text) if charge_text else 0

                total = amt + charge
                total_bill.setText(f"{total:.2f}")
            except ValueError:
                total_bill.setText("0.00")


        #connect signal to handle selection changes
        def on_client_selected(index):
            selected_id = client.itemData(index)
            if selected_id is None:
                present_reading.setEnabled(False)
                present_reading.setStyleSheet(readonly_style)
                return
            IadminPageBack = adminPageBack()
            client_info = IadminPageBack.fetch_client_by_id(selected_id)[0] # get client info by id
            meter_id = client_info[5] # meter id
            client_categ_id = client_info[7] # category id
            previous_reading.setText(str(IadminPageBack.fetch_meter_previous_reading(meter_id))) # get previous reading from meter id # get previous reading from meter id
            # Fetch rate blocks
            self.rate_blocks = IadminPageBack.fetch_rate_blocks_by_categ(client_categ_id)


            # Store categ_id for use in bill creation
            self.categ_id = client_categ_id
            present_reading.setEnabled(True)
            present_reading.setStyleSheet(input_style)
            

        client.currentIndexChanged.connect(on_client_selected)

        def on_present_reading_changed():
            try:
                prev = float(previous_reading.text())
                pres = float(present_reading.text())
                if pres < prev:
                    cubic_meter_consumed.setText("0")
                    amount.setText("0.00")
                    return

                consumed = pres - prev
                cubic_meter_consumed.setText(str(consumed))

                total_amount = 0

                for block in self.rate_blocks:
                    is_minimum = block[1]           # bool: True if it's the minimum block
                    min_c = block[2]                # min cubmic meter for this block (None if is_minimum)
                    max_c = block[3] if block[3] is not None else float('inf') # max cubic meter for this block (None if is_minimum)
                    rate = block[4]                 # rate per cu.m (None if is_nimum)

                    if is_minimum:
                        total_amount += rate

                    elif not is_minimum and consumed >= min_c:
                        applied_volume = max(0, min(consumed, max_c) - min_c + 1)
                        total_amount += applied_volume * rate

                amount.setText(f"{total_amount:.2f}")
                update_total_bill()

            except ValueError:
                cubic_meter_consumed.setText("0")
                amount.setText("0.00")

        present_reading.textChanged.connect(on_present_reading_changed)

        def update_total_charge():
            try:
                sub = float(subscribe_capital.text()) if subscribe_capital.text() else 0
                late = float(late_payment.text()) if late_payment.text() else 0
                pen = float(penalty.text()) if penalty.text() else 0
                total = sub + late + pen
                total_charge.setText(f"{total:.2f}")
                update_total_bill()
            except ValueError:
                total_charge.setText("0.00")
        
        subscribe_capital.textChanged.connect(update_total_charge)
        late_payment.textChanged.connect(update_total_charge)
        penalty.textChanged.connect(update_total_charge)

        def validate_billing_data():
                error_style = """
                    QLineEdit, QDateEdit, QComboBox {
                        padding: 8px;
                        border: 1px solid red;
                        border-radius: 4px;
                        font-family: 'Roboto', sans-serif;
                        min-width: 250px;
                    }
                """
                normal_style = input_style
                errors = []
                has_empty_fields = False
                
                # Reset all styles
                for widget in [client, reading_date, present_reading, due_date, 
                              subscribe_capital, late_payment, penalty]:
                    widget.setStyleSheet(normal_style)
                
                # Check for empty fields first
                if client.currentData() is None:
                    client.setStyleSheet(error_style)
                    has_empty_fields = True
                
                if not present_reading.text().strip():
                    present_reading.setStyleSheet(error_style)
                    has_empty_fields = True
                
                for field in [subscribe_capital, late_payment, penalty]:
                    if not field.text().strip():
                        field.setStyleSheet(error_style)
                        has_empty_fields = True
                
                if has_empty_fields:
                    msg = QtWidgets.QMessageBox(dialog)
                    msg.setWindowTitle("Validation Error")
                    msg.setText("All fields are needed to be filled")
                    msg.setIcon(QtWidgets.QMessageBox.Warning)
                    msg.setStyleSheet("QMessageBox { background-color: white; }")
                    msg.exec_()
                    return False
                
                # Continue with other validations only if no empty fields
                try:
                    prev = float(previous_reading.text() or 0)
                    pres = float(present_reading.text() or 0)
                    if pres <= prev:
                        present_reading.setStyleSheet(error_style)
                        errors.append("\nPresent reading must be greater than previous reading\n")
                except ValueError:
                    present_reading.setStyleSheet(error_style)
                    errors.append("\nInvalid present reading value\n")
                
                if reading_date.date() > QtCore.QDate.currentDate():
                    reading_date.setStyleSheet(error_style)
                    errors.append("\nReading date cannot be in the future\n")
                
                if due_date.date() <= reading_date.date():
                    due_date.setStyleSheet(error_style)
                    errors.append("\nDue date must be after reading date\n")
                
                # Additional charges validation for non-empty fields
                for field, field_name in [(subscribe_capital, "Subscribe Capital"), 
                                        (late_payment, "Late Payment"), 
                                        (penalty, "Penalty")]:
                    try:
                        value = float(field.text())
                        if value < 0:
                            field.setStyleSheet(error_style)
                            errors.append(f"\n{field_name} cannot be negative\n")
                    except ValueError:
                        field.setStyleSheet(error_style)
                        errors.append(f"\nInvalid {field_name} amount\n")
                
                if errors:
                    msg = QtWidgets.QMessageBox(dialog)
                    msg.setWindowTitle("Validation Error")
                    msg.setText("\n\n".join(errors))
                    msg.setIcon(QtWidgets.QMessageBox.Warning)
                    msg.setStyleSheet("QMessageBox { background-color: white; }")
                    msg.exec_()
                    return False
                return True

        def save_bill():
                if not validate_billing_data():
                        return
                try:
                    
                    #backend style
                    #kung kani imo gamiton make sure lang sad nga dapat masave sila tulo dungan walay usa ma fail
                    # akong paabot sat ulo kay ang create reading, create billing, update meter sa iyang reading og last reading date
                    client_id = client.currentData()  # get selected client_id from comboBox
                    prev_read = float(previous_reading.text())
                    pres_read = float(present_reading.text())
                    read_date = reading_date.date().toPyDate()
                    meter_id = IadminPageBack.fetch_client_by_id(client_id)[0][5]  # get meter id from client id

                    reading_id = IadminPageBack.add_reading(read_date, prev_read, pres_read, meter_id) # uncomment ig ready, himog add reading nga function nya e return ang reading id, paki edit nlng pd sa adminback para matest nmo
                    IadminPageBack.update_meter_latest_reading(pres_read, read_date, meter_id) # uncomment sad ig ready, bali maupdate ang last reading sa meter og ang last reading date
                    billing_data = {
                        "billing_due": due_date.date().toPyDate(),
                        "billing_total": float(total_bill.text()) if total_bill.text() else 0,
                        "billing_consumption": float(cubic_meter_consumed.text()) if cubic_meter_consumed.text() else 0,
                        "reading_id": reading_id, # ilisi ang none og reading id kung successfully maka create na
                        "client_id": client_id,
                        "categ_id": self.categ_id,
                        "billing_date": read_date,
                        "billing_status": "TO BE PRINTED",
                        "billing_amount": float(amount.text()) if amount.text() else 0,
                        "billing_sub_capital": float(subscribe_capital.text()) if subscribe_capital.text() else 0,
                        "billing_late_payment": float(late_payment.text()) if late_payment.text() else 0,
                        "billing_penalty": float(penalty.text()) if penalty.text() else 0,
                        "billing_total_charge": float(total_charge.text()) if total_charge.text() else 0
                    }

                    print("READY TO SAVE:", billing_data) # testing rani para check if naget ba ang tanan
                    print(pres_read)
                    IadminPageBack.add_billing(billing_data['billing_due'],
                                            billing_data['billing_total'],
                                            billing_data['billing_consumption'],
                                            billing_data['reading_id'],
                                            billing_data['client_id'],
                                            billing_data['categ_id'],
                                            billing_data['billing_date'],
                                            billing_data['billing_status'],
                                            billing_data['billing_amount'],
                                            billing_data['billing_sub_capital'],
                                            billing_data['billing_late_payment'],
                                            billing_data['billing_penalty'],
                                            billing_data['billing_total_charge'],) # tanggala ang comment kung ready na ang billing repo


                    QtWidgets.QMessageBox.information(dialog, "Success", "Billing information saved successfully.")
                    dialog.accept()
                    updated_data = IadminPageBack.fetch_billing()
                    self.populate_table(updated_data)

                #maka update bisag error
                except Exception as e:
                    QtWidgets.QMessageBox.warning(dialog, "Error", f"Failed to save billing: {str(e)}")

        save_btn.clicked.connect(save_bill)

        dialog.exec_()

class ScrollableTextWidget(QtWidgets.QWidget):
    
    def __init__(self, text, parent=None):
        super(ScrollableTextWidget, self).__init__(parent)
        layout = QtWidgets.QHBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Create scrollable text area
        self.text_area = QtWidgets.QScrollArea()
        self.text_area.setWidgetResizable(True)
        self.text_area.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)  
        self.text_area.setVerticalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
        self.text_area.setFrameShape(QtWidgets.QFrame.NoFrame)
        
        # Create a label with the text
        self.label = QtWidgets.QLabel(text)
        self.label.setTextFormat(QtCore.Qt.PlainText)
        self.label.setAlignment(QtCore.Qt.AlignLeft | QtCore.Qt.AlignVCenter)
        
        # Add label to scroll area
        self.text_area.setWidget(self.label)
        
        # Add scroll area to layout
        layout.addWidget(self.text_area)
        
        # Set the widget's style
        self.setStyleSheet("""
            QScrollArea {
                background-color: transparent;
            }
            QLabel {
                background-color: transparent;
                padding-left: 4px;
                padding-right: 4px;
            }
            QScrollBar:horizontal {
                height: 10px;
                background: transparent;
                margin: 0px;
            }
            QScrollBar::handle:horizontal {
                background: #c0c0c0;
                min-width: 20px;
                border-radius: 5px;
            }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
                width: 0px;
            }
        """)
        
        # Add tooltip for the full text
        self.setToolTip(text)

        # Install event filter to track mouse events
        self.installEventFilter(self)
        
    def text(self):
        return self.label.text()
    
    def eventFilter(self, obj, event):
        if obj is self:
            if event.type() == QtCore.QEvent.Enter:
                # Show scrollbar when mouse enters
                self.text_area.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAsNeeded)
                return True
            elif event.type() == QtCore.QEvent.Leave:
                # Hide scrollbar when mouse leaves
                self.text_area.setHorizontalScrollBarPolicy(QtCore.Qt.ScrollBarAlwaysOff)
                return True
        return super(ScrollableTextWidget, self).eventFilter(obj, event)


class TextEllipsisDelegate(QtWidgets.QStyledItemDelegate):
    
    def __init__(self, parent=None):
        super(TextEllipsisDelegate, self).__init__(parent)
        
    def paint(self, painter, option, index):
        # Use default painting
        super(TextEllipsisDelegate, self).paint(painter, option, index)
        
    def helpEvent(self, event, view, option, index):
        # Show tooltip when hovering if text is truncated
        if not event or not view:
            return False
            
        if event.type() == QtCore.QEvent.ToolTip:
            # Get the cell widget
            cell_widget = view.cellWidget(index.row(), index.column())
            
            if cell_widget and isinstance(cell_widget, ScrollableTextWidget):
                # Show tooltip for ScrollableTextWidget
                QtWidgets.QToolTip.showText(event.globalPos(), cell_widget.text(), view)
                return True
            else:
                # For standard items
                item = view.itemFromIndex(index)
                if item:
                    text = item.text()
                    width = option.rect.width()
                    metrics = QtGui.QFontMetrics(option.font)
                    elidedText = metrics.elidedText(text, QtCore.Qt.ElideRight, width)
                    
                    # If text is truncated, show tooltip
                    if elidedText != text:
                        QtWidgets.QToolTip.showText(event.globalPos(), text, view)
                    else:
                        QtWidgets.QToolTip.hideText()
                    
                    return True
                
        return super(TextEllipsisDelegate, self).helpEvent(event, view, option, index)        
        

          


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = EmployeeBillingPage()
    window.show()
    sys.exit(app.exec_())
